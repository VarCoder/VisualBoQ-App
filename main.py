import os
import shutil
import time
import warnings
from io import BytesIO
from pathlib import Path
from typing import List

import requests
from docx import Document as WordDocument
from docx.shared import Inches
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException
from PIL import Image as PILimage
from selenium import webdriver
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webelement import WebElement
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

# TODO: Implement Color for Infinity

warnings.filterwarnings('ignore')
os.chdir(os.path.dirname(os.path.abspath(__file__)))

designerToWeb = {
    "1 Switch Plate": "1 or 2 Switch Plate",
    "2 Switch Plate": "1 or 2 Switch Plate",
    "3 Switch Plate": "3 Switch Plate",
    "4 Switch Plate": "4 Switch Plate",
    "6 Switch Plate": "6 Switch Plate",
    "8 Switch Plate - HR": "8 Switch - Horizontal",
    "8 Switch Plate - VR": "8 Switch - Vertical",
    "12 Switch Plate": "12 Switch"
}


def xlToWebDict(sheet):
    if sheet == "Infinity":
        switch_types = ['1 Gang', '2 Gang', '3 Gang', '4 Gang', '1 Gang Profile Keypad', '2 Gang Profile Keypad', '3 Gang Profile Keypad', '4 Gang Profile Keypad', '6 Gang Profile Keypad', 'Blinds', '2 Blinds', 'Curtain',
                        '2 Curtain', 'Door Bell', 'Fan Dimmer', 'Light Dimmer', '2 Light Dimmer', '3 Light Dimmer', 'Tunable', 'Socket (5-15 Amps)', 'Socket (2 USB+Switch)', 'C Type', 'Socket with C Type', 'HDMI USB', 'Cable', 'Data', 'Telephone']
        XL_TO_WEB = {item: item for item in switch_types}
        XL_TO_WEB["1 Gang - WR(S)"] = '1 Gang'
        XL_TO_WEB["2 Gang - WR(S)"] = '2 Gang'
        XL_TO_WEB["3 Gang - WR(S)"] = '3 Gang'
        XL_TO_WEB["4 Gang - WR(S)"] = '4 Gang'
        XL_TO_WEB["Light Dimmer (M)"] = 'Light Dimmer'
        XL_TO_WEB["Light Dimmer (S)"] = 'Light Dimmer'
        XL_TO_WEB["T Light Dimmer"] = 'Light Dimmer'
        XL_TO_WEB["T Light Dimmer (M)"] = 'Light Dimmer'
        XL_TO_WEB["T Light Dimmer (S)"] = 'Light Dimmer'
        XL_TO_WEB["1 Gang (M)"] = '1 Gang'
        XL_TO_WEB["2 Gang (M)"] = '2 Gang'
        XL_TO_WEB["3 Gang (M)"] = '3 Gang'
        XL_TO_WEB["4 Gang (M)"] = '4 Gang'
        XL_TO_WEB["Socket (USB+C-type(2A)+Switch)"] = "Socket with C Type"
        XL_TO_WEB["Telephone Socket"] = "Telephone"
        XL_TO_WEB["Cable Socket"] = "Cable"
        XL_TO_WEB["Data Socket"] = "Data"

    else:
        switch_types = ['1 Gang', '2 Gang', '3 Gang', '4 Gang', '1 Gang Profile Keypad', '2 Gang Profile Keypad', '3 Gang Profile Keypad', '4 Gang Profile Keypad', '6 Gang Profile Keypad', 'Blinds', '2 Blinds', 'Curtain',
                        '2 Curtain', 'Door Bell', 'Fan Dimmer', 'Light Dimmer', '2 Light Dimmer', '3 Light Dimmer', 'Tunable', 'Socket (5-15 Amps)', 'Socket (2 USB+Switch)', 'C Type', 'Socket with C Type', 'HDMI USB', 'Cable', 'Data', 'Telephone']

        # switch_types = ['1 Gang', '2 Gang', '3 Gang', '4 Gang', 'Blinds', '2 Blinds', 'Curtain',
        #                 '2 Curtain', 'Door Bell', 'Fan Dimmer', 'Light Dimmer', '2 Light Dimmer', '3 Light Dimmer', 'Tunable', 'Socket (5-15 Amps)', 'Socket (2 USB+Switch)', 'C Type', 'Socket with C Type', 'HDMI USB', 'Cable', 'Data', 'Telephone']
        XL_TO_WEB = {item: item for item in switch_types}
        XL_TO_WEB["2 Gang - WR(S)"] = '2 Gang'
        XL_TO_WEB["1 Gang - WR(S)"] = "1 Gang"
        XL_TO_WEB["2 Gang (M)"] = '2 Gang'
        XL_TO_WEB["1 Gang (M)"] = '1 Gang'
        XL_TO_WEB["Socket (USB+C-type(2A)+Switch)"] = "Socket with C Type"
        # XL_TO_WEB["Socket (USB+C-type(2A)+Switch)"] = "Socket (2 USB+Switch)"
        XL_TO_WEB["HDMI Socket"] = "HDMI"

        XL_TO_WEB["DND Call switch"] = "DND Call switch"
        XL_TO_WEB["Thermostat"] = "Thermostat"
        XL_TO_WEB["Panic Button"] = "Panic Button"
        XL_TO_WEB["Motion Sensor"] = "Motion Sensor"
        XL_TO_WEB["Card Key"] = "Card Key"
        XL_TO_WEB["Dummy + Backbox"] = "Dummy + Backbox"
        XL_TO_WEB["Telephone Socket"] = "Telephone"
        XL_TO_WEB["Foot Lamp"] = "Foot Lamp"

    return XL_TO_WEB


def findExcelFile():
    for file in os.listdir():
        try:
            wb = load_workbook(file, read_only=True)
        except InvalidFileException:
            continue
        else:
            return file
    raise Exception("No File Readable to Excel in this Directory")


def removeFiles(path, dir):
    if path in os.listdir():
        os.remove(path)

    if dir in os.listdir():
        shutil.rmtree(dir)


def setImageDpi(path, dpi):
    image = PILimage.open(path)
    # image.convert("RGB").save(path, dpi=(dpi, dpi))
    image.save(path, dpi=(dpi, dpi))


class Doc():
    def __init__(self, template="assets/template.docx", fileName="Proposal", logo="assets/logo.png"):
        self.doc = WordDocument(template)
        self.fileName = fileName+".docx"
        self.logo = logo

    def addCoverPage(self, clientDetails):

        self.clientTable = self.doc.add_table(rows=1, cols=2, style="NewStyle")
        self.clientTable.rows[0].cells[0].text = "Client Details: "

        for heading, info in clientDetails:
            row = self.clientTable.add_row().cells
            row[0].text = heading
            row[1].text = info

        self.doc.add_page_break()

    def addHeader(self):
        header = self.doc.sections[0].header
        paragraph = header.paragraphs[0]

        logo_run = paragraph.add_run()
        logo_run.add_picture(self.logo, width=Inches(2.5))

    def addRun(self, img, desc1, desc2, desc3, img2=None):
        p = self.doc.add_paragraph()
        r = p.add_run()
        if img2 != None:
            r.add_picture(img2)
            r.add_break()
        r.add_picture(img)
        r.add_break()
        r.add_text(desc1)
        r.add_break()
        r.add_text(desc2)
        r.add_break()
        r.add_text(desc3)

    def save(self):
        self.doc.save(f"{self.fileName}")


def find(it, pred, default=None):
    return next(filter(pred, it), default)


class Sheet():
    # Goal for this sheet is to have all the info we want to extract
    # Column info for space and product descriptions
    # Column info for switch info
    # Potentially more column info for proposal information
    def __init__(self, name, num):
        self.name = name
        self.info = {}
        self.num = num
        self.maxRow = None

    def __repr__(self):
        return f"Sheet({self.name})"

    def addColInfo(self, colStart, info, colEnd=None):
        if colEnd == None:
            colEnd = colStart

        self.info[info] = [
            colStart, colEnd] if colStart != colEnd else colStart


class Agent():

    def __init__(self, wb, dir="tmp", sheets: List[Sheet] = [], url="https://app.smarttouchswitch.com/"):
        self.options = webdriver.ChromeOptions()
        # self.options.add_argument("--disable-gpu")
        self.options.add_argument("--headless")
        self.options.add_argument("--window-size=1920,1080")
        self.options.add_argument('log-level=3')
        self.url = url
        self.driver = webdriver.Chrome(options=self.options)
        self.maxWait = 3

        self.wb = wb
        self.dir = dir
        self.XL_WEB_INF = xlToWebDict("Infinity")
        self.XL_WEB_DES = xlToWebDict("Designer")
        self.startInd = 14

        self.sheetObjs = sheets
        self.sheets = [self.wb[sheet.name] for sheet in self.sheetObjs]
        self.maxRows = [self.getMaxLen(sheet) for sheet in self.sheets]

        self.indArr = [elem if self.maxRows[elem] >=  # Finds any empty sheets and marks them
                       self.startInd else None for elem in range(len(self.sheets))]

        def reformat(x): return [x[elem] for elem in range(
            len(x)) if elem == self.indArr[elem]]

        self.maxRows = reformat(self.maxRows)
        self.sheets = reformat(self.sheets)
        self.sheetObjs = reformat(self.sheetObjs)

        if self.dir in os.listdir():
            os.chdir(Path(self.dir))
        else:
            os.mkdir(Path(self.dir))
            os.chdir(Path(self.dir))

    def openToIndia(self):
        # Waits until page is fully loaded
        self.driver.implicitly_wait(self.maxWait)
        self.driver.get(self.url)

        try:
            button = self.driver.find_element(
                By.CLASS_NAME, "build-action-button")
            self.click(button)  # Opens to the Builder Menu

        except TimeoutException:
            print("Loading took too much time!")

    def getMaxLen(self, sheet):
        # For Infinity and Designer B is the Date
        # For Infinity O is a Color
        # For Designer O is the System
        # These tell how many rows of data there is

        for i in range(self.startInd, sheet.max_row):
            if sheet[f"B{i}"].value == None and sheet[f"O{i}"].value == None:
                return i-1

    def getClientDetails(self):
        sheet = self.wb["Space"]
        col = [str(cell[0].value).replace('\n', '')
               for cell in sheet["A7:A13"]]
        col2 = [str(cell[0].value).replace('\n', '')
                for cell in sheet["C7:C13"]]
        self.clientDetails = list(zip(col, col2))

    def click(self, element: WebElement):
        self.driver.execute_script("arguments[0].click();", element)

    def getModules(self):

        self.modulesFinal = []
        for sheet in range(len(self.sheets)):
            self.modules = []
            start, end = self.sheetObjs[sheet].info["Modules"]
            cell_range = f"{start}{self.startInd}:{end}{self.maxRows[sheet]}"
            maxModuleSize = ord(end) - ord(start) + 1
            cnt = 0
            for column in self.sheets[sheet][cell_range]:
                for cell in column:
                    if cnt % maxModuleSize == 0:
                        self.modules.append([self.sheetObjs[sheet], cell.row])
                    if cell.value != None:
                        self.modules[cnt//maxModuleSize].append(cell.value)
                        # print(self.modules[cnt//maxModuleSize])
                    cnt += 1
            self.modulesFinal += self.modules

        tmpModules = []
        cnt = 0
        for module in self.modulesFinal:
            tmpModules.append([module[:2]])
            for item in module[2:]:
                XL_WEB = self.XL_WEB_INF if module[0].name == "Infinity" else self.XL_WEB_DES
                if item in XL_WEB:
                    tmpModules[cnt].append(XL_WEB[item])
                else:
                    continue
            # if a module is empty (has no valid switches or excel row is empty)
            if len(tmpModules[cnt]) == 1:
                tmpModules.pop()
            else:
                cnt += 1
        """
        self.modules : list[module] ->  
        module = [[Sheet(Infinity or Designer),rowOfInformation], Module 1, Module 2, Module 3]
        """
        self.modules = tmpModules

    def getColors(self):  # sheet is the indice of the sheet
        self.colorsFinal = []
        for sheet in range(len(self.sheets)):
            self.colors = []
            start, end = self.sheetObjs[sheet].info["Colors"]
            cell_range = f"{start}{self.startInd}:{end}{self.maxRows[sheet]}"

            colorArr = ord(end) - ord(start) + 1
            cnt = 0
            for column in self.sheets[sheet][cell_range]:
                for cell in column:
                    if cnt % colorArr == 0:
                        self.colors.append(
                            [[self.sheetObjs[sheet], cell.row]]
                        )
                    if cell.value != None:
                        self.colors[cnt//colorArr].append(cell.value)
                    cnt += 1
            self.colorsFinal += self.colors
        self.colors = self.colorsFinal
        # else:
        #     start, end = self.sheetObjs[sheet].info
        # print(self.colors)

    def clickColor(self, level: str, colorProfile: str, colorInfo: List):
        try:
            WebDriverWait(self.driver, self.maxWait).until(
                EC.visibility_of_any_elements_located(
                    (By.CLASS_NAME, "mod-label"))
            )
        except:
            self.click(self.colorPanel)
            # self.driver.execute_script(
            #     "arguments[0].click();", self.colorPanel)
            WebDriverWait(self.driver, self.maxWait).until(
                EC.visibility_of_any_elements_located(
                    (By.CLASS_NAME, "mod-label"))
            )

        #colorProfile = [[Sheet(Designer), row], OuterGlass, OuterFrame, InnerGlass, InnerFrame]
        colorType = self.driver.find_element(
            By.XPATH, f"//span[text()=\'{level}\']"
        )
        # colorType.click()
        self.click(colorType)
        if level == "Outer Surface":
            WebDriverWait(self.driver, self.maxWait).until(
                EC.visibility_of_any_elements_located(
                    (By.CLASS_NAME, "fab-label-inner"))
            )
            glass = self.driver.find_element(
                By.CLASS_NAME, "fab-label-inner"
            )
            self.click(glass)
        color = self.driver.find_element(
            By.XPATH, f"//a[@glass-color=\'{colorProfile.lower()}\'][@data-colortype='glass']"
        )
        self.click(color)

    def clickModules(self):
        for moduleInd in range(len(self.modules)):
            start = time.perf_counter()
            modules = self.modules[moduleInd]
            self.openToIndia()

            WebDriverWait(self.driver, self.maxWait).until(
                EC.visibility_of_any_elements_located(
                    (By.CLASS_NAME, "mod-label"))
            )

            if modules[0][0].name == "Designer":
                ws = self.wb["Designer"]
                switchPlate = designerToWeb[ws[f"{modules[0][0].info['Switch']}{modules[0][1]}"].value]
                modType = self.driver.find_element(
                    By.XPATH, f"//span[text()=\'{switchPlate}\']")
                self.click(modType)
            else:
                # TODO: Implement vertical switches (when excel is updated)
                modType = self.driver.find_elements(
                    By.CLASS_NAME, "mod-label")[len(modules)-2]
                self.click(modType)

            sizePanel = self.driver.find_element(
                By.CSS_SELECTOR, 'div[data-panelid=".sizePanel"]'
            )
            self.click(sizePanel)
            self.driver.implicitly_wait(1)
            WebDriverWait(self.driver, 5).until(
                EC.invisibility_of_element(
                    (By.CLASS_NAME, "mod-label")
                )
            )
            modPanel = self.driver.find_element(
                By.CSS_SELECTOR, 'div[data-panelid=".modulePanel"]'
            )
            self.click(modPanel)
            
            WebDriverWait(self.driver, 5).until(
                EC.visibility_of_any_elements_located(
                    (By.CLASS_NAME, "module-type-label"))
            )

            # Skip the module info and get the modules themselves
            for module in modules[1:]:
                self.driver.implicitly_wait(20)
                modToClick = self.driver.find_element(
                    By.XPATH, f"//div[text()=\'{module}\']"
                )
                self.click(modToClick)
            self.driver.implicitly_wait(1)
            WebDriverWait(self.driver, self.maxWait).until(
                EC.invisibility_of_element(
                    (By.CLASS_NAME, "module-type-label")
                )
            )
            self.colorPanel = WebDriverWait(self.driver, self.maxWait).until(
                lambda driver: driver.find_element(
                    By.CSS_SELECTOR, 'div[data-panelid=".colorPanel"]'
                )
            )
            self.click(self.colorPanel)
            colorProfile = find(
                self.colors, lambda x: x[0] == modules[0]
            )
            # #TODO: Implement Colors
            colorInfo = colorProfile[0]
            colorProfile = colorProfile[1:]
            if "TBD" in colorProfile:
                self.click(self.colorPanel)
                self.screenshot(index=moduleInd)
                end = time.perf_counter()
                print(f"{end-start} Seconds taken for the switch")
                continue
            response = requests.get(
                f"https://app.smarttouchswitch.com/modules/components/images/frames/{colorProfile[0]}{colorProfile[1]}-Frame.png"
            )
            frame = PILimage.open(BytesIO(response.content))
            framePath = f"frame_{moduleInd}.png"
            frame.save(framePath)
            modules[0].append(framePath)
            if modules[0][0].name == "Designer":
                self.clickColor(
                    "Outer Surface", colorProfile=colorProfile[0], colorInfo=colorInfo)
                self.clickColor(
                    "Outer Frame", colorProfile=colorProfile[1], colorInfo=colorInfo)
                self.clickColor(
                    "Inner Surface", colorProfile=colorProfile[2], colorInfo=colorInfo)
                self.clickColor(
                    "Inner Frame", colorProfile=colorProfile[3], colorInfo=colorInfo)
            else:
                self.clickColor(
                    "Outer Surface", colorProfile=colorProfile[0], colorInfo=colorInfo)
                self.clickColor(
                    "Outer Frame", colorProfile=colorProfile[1], colorInfo=colorInfo)
            self.click(self.colorPanel)
            self.screenshot(index=moduleInd)
            end = time.perf_counter()
            print(f"{end-start} Seconds taken for the switch")

    def screenshot(self, index):
        final_switch = self.driver.find_element(
            By.CLASS_NAME, "switch-panel")
        final_switch.screenshot(f'switch_{index}.png')
        print(f"Switch {index} Completed")

    def close(self):
        self.driver.quit()
        os.chdir("..")

    def getCol(self, sheet, colLetter):
        colVals = []
        for column in self.wb[self.sheetObjs[sheet].name][f"{colLetter}{self.startInd}:{colLetter}{self.maxRows[sheet]}"]:
            for cell in column:
                colVals.append(cell.value)
        return colVals

    def publish(self, fileName="Proposal", debug=False):
        self.getClientDetails()
        document = Doc(fileName=fileName)
        self.docx = document.fileName
        document.addCoverPage(self.clientDetails)
        document.addHeader()
        if debug:
            for img_path in os.listdir(self.dir):
                self.switchPath = str(Path(os.path.join(self.dir, img_path)).absolute())
                spaceText = "Space: "
                prodText = "Product Description: "
                setImageDpi(self.switchPath, 96*2)

                document.addRun(self.switchPath, spaceText, prodText)
            document.save()
        else:
            for switch in range(len(self.modules)):
                sheetObj = self.modules[switch][0][0]
                row = self.modules[switch][0][1]
                frameImg = None
                if isinstance(self.modules[switch][0][-1], str) and self.modules[switch][0][-1].endswith(".png"):
                    frameImg = str(
                        Path(os.path.join(self.dir, self.modules[switch][0][-1])).absolute())
                    setImageDpi(frameImg, 96*2)
                sheet = self.wb[sheetObj.name]

                prodDesc = "Product Description: " + \
                    str(sheet[f"{sheetObj.info['Product']}{row}"].value)
                space = "Space: " + \
                    str(sheet[f"{sheetObj.info['Space']}{row}"].value)
                prodType = "Product Type : " + \
                    str(sheetObj.name)
                self.switchPath = str(
                    Path(os.path.join(self.dir, f"switch_{switch}.png")).absolute())
                setImageDpi(self.switchPath, 96*2)
                document.addRun(self.switchPath, space, prodType, prodDesc, frameImg)
            document.save()


